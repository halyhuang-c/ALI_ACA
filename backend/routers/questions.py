from typing import Optional

from fastapi import APIRouter, Depends, Query, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy import or_, cast, String
from sqlalchemy.orm import Session, joinedload

from database import get_db
from models import Question, Answer, QuestionTag, Tag
from schemas import QuestionOut, QuestionSearchResult

router = APIRouter()


def _to_question_out(q: Question, duplicate_info: Optional[dict] = None) -> QuestionOut:
    answer: Optional[Answer] = q.answer
    raw = answer.raw_response if answer else None
    history = []
    if isinstance(raw, dict):
        history = raw.get("history") or []
    return QuestionOut(
        id=q.id,
        image_id=q.image_id,
        image_filename=q.image.filename if q.image else None,
        question_text=q.question_text,
        options=q.options,
        question_type=q.question_type,
        correct_answer=q.correct_answer,
        category=q.category,
        subcategory=q.subcategory,
        norm_text=q.norm_text,
        dedup_hash=q.dedup_hash,
        is_duplicate=q.is_duplicate,
        duplicate_of_id=q.duplicate_of_id,
        created_at=q.created_at,
        answer_text=answer.answer if answer else None,
        explanation=answer.explanation if answer else None,
        tags=answer.tags if answer else None,
        is_correct=answer.is_correct if answer else None,
        answer_id=answer.id if answer else None,
        answer_model=answer.model if answer else None,
        review_status=answer.review_status if answer else None,
        is_reanswered=bool(answer.is_reanswered) if answer else False,
        answer_history=history if isinstance(history, list) else [],
        duplicate_ids=duplicate_info.get("ids", []) if duplicate_info else None,
        duplicate_answer_conflict=duplicate_info.get("answer_conflict", False) if duplicate_info else None,
    )


def _paginate_question_ids(base, page: int, page_size: int):
    id_query = base.with_entities(Question.id).distinct().order_by(Question.id.asc())
    total = id_query.count()
    skip = (page - 1) * page_size
    rows = id_query.offset(skip).limit(page_size).all()
    ids = [r[0] for r in rows]
    return ids, total


@router.get("/questions", response_model=list[QuestionOut])
def list_questions(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    tag: Optional[str] = Query(None),
    category: Optional[str] = Query(None),
    subcategory: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    base = db.query(Question).filter(Question.is_duplicate.is_(False))
    if tag:
        norm_tag = tag.strip().lower()
        base = (
            base.join(QuestionTag, QuestionTag.question_id == Question.id)
            .join(Tag, Tag.id == QuestionTag.tag_id)
            .filter(Tag.name == norm_tag)
        )
    if category:
        base = base.filter(Question.category == category.strip())
    if subcategory:
        base = base.filter(Question.subcategory == subcategory.strip())

    ids, _ = _paginate_question_ids(base, page, page_size)
    if not ids:
        return []
    questions = (
        db.query(Question)
        .options(joinedload(Question.image), joinedload(Question.answer))
        .filter(Question.id.in_(ids))
        .order_by(Question.id.asc())
        .all()
    )
    return [_to_question_out(q) for q in questions]


@router.get("/questions/search", response_model=QuestionSearchResult)
def search_questions(
    keyword: Optional[str] = Query(None),
    tag: Optional[str] = Query(None),
    category: Optional[str] = Query(None),
    subcategory: Optional[str] = Query(None),
    only_wrong: bool = Query(False),
    exclude_wrong: bool = Query(False),
    question_type: Optional[str] = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    db: Session = Depends(get_db),
):
    base = db.query(Question).filter(Question.is_duplicate.is_(False))

    if tag:
        norm_tag = tag.strip().lower()
        base = (
            base.join(QuestionTag, QuestionTag.question_id == Question.id)
            .join(Tag, Tag.id == QuestionTag.tag_id)
            .filter(Tag.name == norm_tag)
        )

    if category:
        base = base.filter(Question.category == category.strip())
    if subcategory:
        base = base.filter(Question.subcategory == subcategory.strip())

    if question_type:
        base = base.filter(Question.question_type == question_type.strip())

    need_answer_join = bool(keyword or only_wrong or exclude_wrong)
    if need_answer_join:
        base = base.outerjoin(Answer, Answer.question_id == Question.id)

    if keyword:
        kw = f"%{keyword}%"
        base = base.filter(
            or_(
                Question.question_text.like(kw),
                cast(Question.options, String).like(kw),
                Answer.answer.like(kw),
                Answer.explanation.like(kw),
            )
        )

    if only_wrong:
        base = base.filter(Answer.is_correct.is_(False))

    if exclude_wrong:
        base = base.filter(or_(Answer.is_correct.is_(True), Answer.is_correct.is_(None)))

    ids, total = _paginate_question_ids(base, page, page_size)
    items: list[QuestionOut] = []
    if ids:
        questions = (
            db.query(Question)
            .options(joinedload(Question.image), joinedload(Question.answer))
            .filter(Question.id.in_(ids))
            .order_by(Question.id.asc())
            .all()
        )
        # 查询这些题目的重复题目信息
        dup_questions = (
            db.query(Question.id, Question.duplicate_of_id, Question.correct_answer)
            .filter(Question.is_duplicate.is_(True), Question.duplicate_of_id.in_(ids))
            .all()
        )
        dup_map: dict[int, list[dict]] = {}
        for dup_id, parent_id, dup_answer in dup_questions:
            dup_map.setdefault(parent_id, []).append({"id": dup_id, "correct_answer": dup_answer})

        for q in questions:
            dups = dup_map.get(q.id, [])
            dup_ids = [d["id"] for d in dups]
            answer_conflict = any(
                d["correct_answer"] and q.correct_answer
                and d["correct_answer"].strip().upper() != q.correct_answer.strip().upper()
                for d in dups
            )
            dup_info = {"ids": dup_ids, "answer_conflict": answer_conflict} if dup_ids else None
            items.append(_to_question_out(q, dup_info))

    return QuestionSearchResult(items=items, total=total, page=page, page_size=page_size)


@router.get("/questions/export-word")
def export_questions_word(
    keyword: Optional[str] = Query(None),
    tag: Optional[str] = Query(None),
    category: Optional[str] = Query(None),
    subcategory: Optional[str] = Query(None),
    only_wrong: bool = Query(False),
    exclude_wrong: bool = Query(False),
    question_type: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    import traceback as _tb
    try:
        return _export_questions_word_impl(keyword, tag, category, subcategory, only_wrong, exclude_wrong, question_type, db)
    except Exception as e:
        _tb.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


def _export_questions_word_impl(keyword, tag, category, subcategory, only_wrong, exclude_wrong, question_type, db):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    # 复用搜索逻辑，不分页，取全部
    base = db.query(Question).filter(Question.is_duplicate.is_(False))

    if tag:
        norm_tag = tag.strip().lower()
        base = (
            base.join(QuestionTag, QuestionTag.question_id == Question.id)
            .join(Tag, Tag.id == QuestionTag.tag_id)
            .filter(Tag.name == norm_tag)
        )

    if category:
        base = base.filter(Question.category == category.strip())
    if subcategory:
        base = base.filter(Question.subcategory == subcategory.strip())

    if question_type:
        base = base.filter(Question.question_type == question_type.strip())

    need_answer_join = bool(keyword or only_wrong or exclude_wrong)
    if need_answer_join:
        base = base.outerjoin(Answer, Answer.question_id == Question.id)

    if keyword:
        kw = f"%{keyword}%"
        base = base.filter(
            or_(
                Question.question_text.like(kw),
                cast(Question.options, String).like(kw),
                Answer.answer.like(kw),
                Answer.explanation.like(kw),
            )
        )

    if only_wrong:
        base = base.filter(Answer.is_correct.is_(False))

    if exclude_wrong:
        base = base.filter(or_(Answer.is_correct.is_(True), Answer.is_correct.is_(None)))

    # 先获取符合条件的 Question ID，再用独立查询加载完整数据（避免 join + joinedload 冲突）
    id_rows = base.with_entities(Question.id).order_by(Question.id.asc()).all()
    id_list = [row[0] for row in id_rows]

    questions = []
    if id_list:
        questions = (
            db.query(Question)
            .options(joinedload(Question.answer))
            .filter(Question.id.in_(id_list))
            .order_by(Question.id.asc())
            .all()
        )

    doc = Document()

    # 设置默认字体（西文 + 东亚）
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    from docx.oxml.ns import qn
    from lxml import etree
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题
    title = doc.add_heading("题目导出", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"共 {len(questions)} 道题目")
    doc.add_paragraph("")

    ORDER = ["A", "B", "C", "D", "E", "F"]

    for idx, q in enumerate(questions, 1):
        answer = q.answer
        is_wrong = answer and answer.is_correct is False

        # 题号 + 题型
        q_type = q.question_type or "未知"
        heading = doc.add_heading(f"{idx}. 【{q_type}】", level=2)

        # 题目正文
        p = doc.add_paragraph(q.question_text or "（无题干）")
        p.paragraph_format.space_after = Pt(4)

        # 选项
        if q.options and isinstance(q.options, dict):
            for key in sorted(q.options.keys(), key=lambda k: ORDER.index(k) if k in ORDER else 99):
                val = q.options.get(key, "")
                p = doc.add_paragraph(f"{key}. {val}", style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)

        # 答案
        if answer:
            if is_wrong:
                # AI 答案与标准答案不一致：显示标准答案 + AI答案和解析
                p = doc.add_paragraph()
                run = p.add_run("标准答案：")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)
                run2 = p.add_run(q.correct_answer or "（未标注）")
                run2.font.color.rgb = RGBColor(0, 128, 0)

                p2 = doc.add_paragraph()
                run3 = p2.add_run(f"AI 答案（{answer.model or '未知'}）：")
                run3.bold = True
                run3.font.color.rgb = RGBColor(200, 0, 0)
                run4 = p2.add_run(answer.answer or "（空）")
                run4.font.color.rgb = RGBColor(200, 0, 0)

                if answer.explanation:
                    p3 = doc.add_paragraph()
                    run5 = p3.add_run("AI 解析：")
                    run5.bold = True
                    p3.add_run(answer.explanation)
            else:
                # 一致或无标准答案：显示答案和解析
                p = doc.add_paragraph()
                run = p.add_run("答案：")
                run.bold = True
                p.add_run(answer.answer or "（空）")

                if answer.explanation:
                    p2 = doc.add_paragraph()
                    run2 = p2.add_run("解析：")
                    run2.bold = True
                    p2.add_run(answer.explanation)
        else:
            if q.correct_answer:
                p = doc.add_paragraph()
                run = p.add_run("标准答案：")
                run.bold = True
                p.add_run(q.correct_answer)

        # 分隔
        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from datetime import datetime
    from urllib.parse import quote
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    ascii_name = f"questions_{ts}.docx"
    cn_name = quote(f"题目导出_{ts}.docx")

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{cn_name}",
        },
    )


@router.get("/questions/by-ids", response_model=list[QuestionOut])
def get_questions_by_ids(
    ids: str = Query(..., description="逗号分隔的题目ID"),
    db: Session = Depends(get_db),
):
    id_list = [int(x.strip()) for x in ids.split(",") if x.strip().isdigit()]
    if not id_list:
        return []
    questions = (
        db.query(Question)
        .options(joinedload(Question.image), joinedload(Question.answer))
        .filter(Question.id.in_(id_list))
        .all()
    )
    q_map = {q.id: q for q in questions}
    return [_to_question_out(q_map[i]) for i in id_list if i in q_map]


@router.get("/questions/categories")
def list_categories(db: Session = Depends(get_db)):
    from sqlalchemy import func
    rows = (
        db.query(Question.category, Question.subcategory, func.count(Question.id))
        .filter(Question.is_duplicate.is_(False))
        .group_by(Question.category, Question.subcategory)
        .all()
    )
    tree: dict[str, dict] = {}
    for category, subcategory, cnt in rows:
        cat = category or "未分类"
        sub = subcategory or "未分类"
        node = tree.setdefault(cat, {"count": 0, "subcategories": {}})
        node["count"] += int(cnt or 0)
        node["subcategories"][sub] = node["subcategories"].get(sub, 0) + int(cnt or 0)
    return {
        "categories": [
            {
                "name": cat,
                "count": info["count"],
                "subcategories": [
                    {"name": s, "count": c} for s, c in sorted(info["subcategories"].items())
                ],
            }
            for cat, info in sorted(tree.items())
        ]
    }


@router.get("/questions/dedup")
def dedup_detail(
    only_duplicates: bool = Query(False),
    keyword: Optional[str] = Query(None),
    db: Session = Depends(get_db),
):
    from services.dedup_service import normalize_text, similarity
    query = db.query(Question)
    if only_duplicates:
        query = query.filter(Question.is_duplicate.is_(True))
    if keyword:
        kw = f"%{keyword}%"
        query = query.filter(
            or_(
                Question.question_text.like(kw),
                Question.norm_text.like(kw),
            )
        )
    questions = query.order_by(Question.dedup_hash.asc(), Question.id.asc()).all()

    groups_map: dict[str, dict] = {}
    for q in questions:
        h = q.dedup_hash or f"__no_hash_{q.id}__"
        norm = q.norm_text if q.norm_text is not None else normalize_text(q.question_text)
        item = {
            "id": q.id,
            "question_text": q.question_text,
            "options": q.options,
            "norm_text": norm,
            "is_duplicate": bool(q.is_duplicate),
            "duplicate_of_id": q.duplicate_of_id,
            "image_id": q.image_id,
        }
        node = groups_map.get(h)
        if node is None:
            groups_map[h] = {
                "dedup_hash": q.dedup_hash or "",
                "norm_text": norm,
                "primary_id": q.duplicate_of_id if q.is_duplicate else q.id,
                "members": [item],
                "is_fuzzy": False,
            }
        else:
            node["members"].append(item)
            # 如果组内有不同的 norm_text，说明是模糊匹配合并的
            if norm != node["norm_text"] and not node["is_fuzzy"]:
                node["is_fuzzy"] = True

    # 对模糊合并的组，计算成员间的相似度
    for g in groups_map.values():
        if g["is_fuzzy"] and len(g["members"]) > 1:
            primary_norm = g["members"][0]["norm_text"]
            for m in g["members"][1:]:
                m["similarity"] = round(similarity(primary_norm, m["norm_text"]), 4)

    groups = sorted(groups_map.values(), key=lambda g: len(g["members"]), reverse=True)

    total_questions = db.query(Question).count()
    total_unique = db.query(Question).filter(Question.is_duplicate.is_(False)).count()
    total_duplicates = total_questions - total_unique
    return {
        "summary": {
            "total_questions": total_questions,
            "total_unique": total_unique,
            "total_duplicates": total_duplicates,
            "dedup_rate": round(total_duplicates / total_questions, 4) if total_questions else 0,
            "group_count": len(groups),
        },
        "groups": groups,
    }


def _compare_answer(ai_ans, correct) -> Optional[bool]:
    a = (ai_ans or "").strip().upper()
    c = (correct or "").strip().upper()
    if not c:
        return None
    if not a:
        return False
    norm_a = sorted(set(a.replace(" ", "").replace(",", "").replace("，", "")))
    norm_c = sorted(set(c.replace(" ", "").replace(",", "").replace("，", "")))
    return norm_a == norm_c


def _normalize_category(value):
    if not value:
        return None
    return str(value).strip() or None


def _normalize_subcategory(value, category):
    if not value:
        return None
    v = str(value).strip()
    if not v or v == category:
        return None
    return v


class ReAnswerRequest(BaseModel):
    config_id: int
    model: str


class BatchReAnswerRequest(BaseModel):
    question_ids: list[int]
    config_id: int
    model: str
    only_wrong: bool = False
    tag: Optional[str] = None
    category: Optional[str] = None


class ReviewRequest(BaseModel):
    answer: Optional[str] = None
    explanation: Optional[str] = None
    is_correct: Optional[bool] = None
    review_status: Optional[str] = None
    adopt_history_index: Optional[int] = None


@router.post("/questions/{question_id}/reanswer", response_model=QuestionOut)
def reanswer_question(
    question_id: int,
    payload: ReAnswerRequest,
    db: Session = Depends(get_db),
):
    from models import LLMConfig
    import llm_client
    import setting_service
    from services.answer_service import normalize_category, normalize_subcategory

    q = (
        db.query(Question)
        .options(joinedload(Question.image), joinedload(Question.answer))
        .filter(Question.id == question_id)
        .first()
    )
    if not q:
        raise HTTPException(status_code=404, detail="题目不存在")
    if q.is_duplicate:
        raise HTTPException(status_code=400, detail="重复题目不支持重解")

    cfg = db.query(LLMConfig).filter(LLMConfig.id == payload.config_id).first()
    if not cfg:
        raise HTTPException(status_code=404, detail="LLM 配置不存在")

    sys_prompt = setting_service.get_setting(db, "answer_system_prompt") or ""
    user_tpl = setting_service.get_setting(db, "answer_prompt_template") or ""

    try:
        result, raw_content = llm_client.answer_question(
            q.question_text,
            q.options,
            cfg.base_url,
            cfg.api_key,
            payload.model,
            sys_prompt,
            user_tpl,
            q.question_type or "",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"模型调用失败：{e}")

    answer = q.answer
    old_snapshot = None
    if answer:
        old_snapshot = {
            "answer": answer.answer,
            "explanation": answer.explanation,
            "model": answer.model,
            "is_correct": answer.is_correct,
            "ts": (answer.created_at.isoformat() if answer.created_at else None),
        }

    ai_ans = result.get("answer")
    correct = q.correct_answer
    is_correct = _compare_answer(ai_ans, correct)
    category = normalize_category(result.get("category"))
    subcategory = normalize_subcategory(result.get("subcategory"), category)
    if category:
        q.category = category
    if subcategory:
        q.subcategory = subcategory

    if answer is None:
        answer = Answer(question_id=q.id)
        db.add(answer)

    raw = dict(answer.raw_response) if isinstance(answer.raw_response, dict) else {}
    history = list(raw.get("history") or [])
    if old_snapshot:
        history.append(old_snapshot)
    raw["history"] = history
    raw["raw"] = raw_content
    raw["parsed"] = result

    answer.answer = ai_ans
    answer.explanation = result.get("explanation")
    answer.tags = result.get("tags") or []
    answer.raw_response = None  # 先置空，强制 SQLAlchemy 检测到变更
    answer.raw_response = raw
    answer.model = payload.model
    answer.is_correct = is_correct
    answer.review_status = "approved" if is_correct is True else "pending"
    answer.is_reanswered = True
    db.commit()
    db.refresh(q)
    return _to_question_out(q)


@router.put("/answers/{answer_id}/review", response_model=QuestionOut)
def review_answer(
    answer_id: int,
    payload: ReviewRequest,
    db: Session = Depends(get_db),
):
    answer = db.query(Answer).filter(Answer.id == answer_id).first()
    if not answer:
        raise HTTPException(status_code=404, detail="答案不存在")

    raw = dict(answer.raw_response) if isinstance(answer.raw_response, dict) else {}
    history = list(raw.get("history") or [])

    if payload.review_status == "approved":
        # 采纳：用 history 中最后一个重解记录覆盖当前答案
        reanswer_entry = None
        reanswer_idx = None
        for i in range(len(history) - 1, -1, -1):
            if history[i].get("is_reanswer"):
                reanswer_entry = history[i]
                reanswer_idx = i
                break
        if reanswer_entry:
            answer.answer = reanswer_entry.get("answer")
            answer.explanation = reanswer_entry.get("explanation")
            answer.model = reanswer_entry.get("model")
            answer.is_correct = reanswer_entry.get("is_correct")
            # 采纳后从 history 中移除该重解记录
            del history[reanswer_idx]
            raw["history"] = history
            answer.raw_response = None
            answer.raw_response = raw
        answer.review_status = "approved"
        answer.is_reanswered = False

    elif payload.review_status == "rejected":
        # 驳回：删除 history 中的重解记录，保留当前答案
        history = [h for h in history if not h.get("is_reanswer")]
        raw["history"] = history
        answer.raw_response = None
        answer.raw_response = raw
        answer.review_status = "rejected"
        answer.is_reanswered = False

    else:
        # 其他状态（pending 等）
        if payload.review_status is not None:
            if payload.review_status not in ("pending", "approved", "rejected"):
                raise HTTPException(status_code=400, detail="非法审核状态")
            answer.review_status = payload.review_status

    # 采纳历史版本（adopt_history_index）
    if payload.adopt_history_index is not None:
        if 0 <= payload.adopt_history_index < len(history):
            chosen = history[payload.adopt_history_index]
            answer.answer = chosen.get("answer")
            answer.explanation = chosen.get("explanation")
            answer.model = chosen.get("model")
            answer.is_correct = chosen.get("is_correct")

    # 手动编辑覆盖
    if payload.answer is not None:
        answer.answer = payload.answer
    if payload.explanation is not None:
        answer.explanation = payload.explanation
    if payload.is_correct is not None:
        answer.is_correct = payload.is_correct

    if payload.answer is not None or payload.is_correct is not None:
        from services.answer_service import _compare_answer as _cmp
        answer.is_correct = _cmp(answer.answer, answer.question.correct_answer)

    db.commit()
    q = (
        db.query(Question)
        .options(joinedload(Question.image), joinedload(Question.answer))
        .filter(Question.id == answer.question_id)
        .first()
    )
    return _to_question_out(q)


def _apply_reanswer_to_question(
    db: Session,
    q: Question,
    cfg,
    model: str,
    sys_prompt: str,
    user_tpl: str,
) -> dict:
    import llm_client
    from services.answer_service import normalize_category, normalize_subcategory

    result, raw_content = llm_client.answer_question(
        q.question_text,
        q.options,
        cfg.base_url,
        cfg.api_key,
        model,
        sys_prompt,
        user_tpl,
        q.question_type or "",
    )

    answer = q.answer
    reanswer_snapshot = None
    if answer:
        # 新答案存入 history，当前答案不变
        reanswer_snapshot = {
            "answer": ai_ans,
            "explanation": result.get("explanation"),
            "model": model,
            "is_correct": is_correct,
            "ts": datetime.utcnow().isoformat(),
            "is_reanswer": True,
        }

    category = normalize_category(result.get("category"))
    subcategory = normalize_subcategory(result.get("subcategory"), category)
    if category:
        q.category = category
    if subcategory:
        q.subcategory = subcategory

    if answer is None:
        answer = Answer(question_id=q.id)
        db.add(answer)

    raw = dict(answer.raw_response) if isinstance(answer.raw_response, dict) else {}
    history = list(raw.get("history") or [])
    if reanswer_snapshot:
        history.append(reanswer_snapshot)
    raw["history"] = history
    raw["raw"] = raw_content
    raw["parsed"] = result

    # 当前答案保持不变，只更新 raw_response 和标记
    answer.raw_response = None
    answer.raw_response = raw
    answer.review_status = "approved" if is_correct is True else "pending"
    answer.is_reanswered = True
    return _to_question_out(q).__dict__


@router.post("/questions/batch-reanswer")
def batch_reanswer(
    payload: BatchReAnswerRequest,
    db: Session = Depends(get_db),
):
    from models import LLMConfig
    import setting_service

    cfg = db.query(LLMConfig).filter(LLMConfig.id == payload.config_id).first()
    if not cfg:
        raise HTTPException(status_code=404, detail="LLM 配置不存在")

    sys_prompt = setting_service.get_setting(db, "answer_system_prompt") or ""
    user_tpl = setting_service.get_setting(db, "answer_prompt_template") or ""

    if payload.question_ids:
        ids = list(payload.question_ids)
    else:
        base = db.query(Question).filter(Question.is_duplicate.is_(False))
        if payload.only_wrong:
            base = (
                base.outerjoin(Answer, Answer.question_id == Question.id)
                .filter(Answer.is_correct.is_(False))
            )
        if payload.tag:
            norm_tag = payload.tag.strip().lower()
            base = (
                base.join(QuestionTag, QuestionTag.question_id == Question.id)
                .join(Tag, Tag.id == QuestionTag.tag_id)
                .filter(Tag.name == norm_tag)
            )
        if payload.category:
            base = base.filter(Question.category == payload.category.strip())
        ids = [row[0] for row in base.with_entities(Question.id).distinct().all()]

    if not ids:
        return {"total": 0, "success": 0, "failed": 0, "items": []}

    items = []
    success = 0
    failed = 0
    for qid in ids:
        q = (
            db.query(Question)
            .options(joinedload(Question.image), joinedload(Question.answer))
            .filter(Question.id == qid)
            .first()
        )
        if not q or q.is_duplicate:
            failed += 1
            items.append({"id": qid, "status": "skipped", "error": "题目不存在或重复"})
            continue
        try:
            out = _apply_reanswer_to_question(db, q, cfg, payload.model, sys_prompt, user_tpl)
            db.commit()
            success += 1
            items.append({"id": qid, "status": "ok", "question": out})
        except Exception as e:
            db.rollback()
            failed += 1
            items.append({"id": qid, "status": "failed", "error": str(e)[:200]})

    return {"total": len(ids), "success": success, "failed": failed, "items": items}


@router.delete("/answers/{answer_id}/history/{index}")
def delete_history_item(
    answer_id: int,
    index: int,
    db: Session = Depends(get_db),
):
    answer = db.query(Answer).filter(Answer.id == answer_id).first()
    if not answer:
        raise HTTPException(status_code=404, detail="答案不存在")
    raw = answer.raw_response if isinstance(answer.raw_response, dict) else {}
    history = raw.get("history") or []
    if index < 0 or index >= len(history):
        raise HTTPException(status_code=400, detail="历史版本索引无效")
    del history[index]
    raw["history"] = history
    answer.raw_response = raw
    db.commit()
    return {"status": "ok", "remaining": len(history)}


@router.delete("/answers/{answer_id}/history")
def clear_history(answer_id: int, db: Session = Depends(get_db)):
    answer = db.query(Answer).filter(Answer.id == answer_id).first()
    if not answer:
        raise HTTPException(status_code=404, detail="答案不存在")
    raw = answer.raw_response if isinstance(answer.raw_response, dict) else {}
    raw["history"] = []
    answer.raw_response = raw
    db.commit()
    return {"status": "ok", "remaining": 0}
